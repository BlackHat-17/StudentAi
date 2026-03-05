from google import genai
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import os
from dotenv import load_dotenv
import os 
load_dotenv()

# -------------------------
# CONFIGURATION
# -------------------------

pdf_file = "test2.pdf"

client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))


# -------------------------
# TEXT CHUNKING
# -------------------------

def chunk_text(text, chunk_size=400):

    words = text.split()
    chunks = []

    for i in range(0, len(words), chunk_size):

        chunk = " ".join(words[i:i + chunk_size])

        chunks.append(chunk)

    return chunks


# -------------------------
# GEMINI OCR FUNCTION
# -------------------------

def gemini_ocr(image_path):

    with Image.open(image_path) as image:

        prompt = """
        Extract all handwritten or printed text from this document page.
        Preserve headings and sentences.
        Ignore drawings unless they contain labels.
        Return only the extracted text.
        """

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt, image]
        )

        return response.text


# -------------------------
# CONVERT PDF TO IMAGES
# -------------------------

print("Converting PDF pages to images...")

images = convert_from_path(pdf_file)

doc = Document()
doc.add_heading("OCR Extracted Document", 0)

all_text = ""


# -------------------------
# PROCESS EACH PAGE
# -------------------------

for i, img in enumerate(images):

    image_path = f"page_{i}.png"

    img.save(image_path)

    try:

        print(f"Reading page {i+1}...")

        text = gemini_ocr(image_path)

        all_text += text + "\n"

        doc.add_heading(f"Page {i+1}", level=1)
        doc.add_paragraph(text)

    except Exception as e:

        print("Error processing page", i+1, e)

    finally:

        if os.path.exists(image_path):
            os.remove(image_path)


# -------------------------
# SAVE WORD DOCUMENT
# -------------------------

doc.save("doc1_output.docx")

print("Word document created: doc1_output.docx")


# -------------------------
# CHECK TEXT EXTRACTION
# -------------------------

if len(all_text.strip()) == 0:

    print("No text extracted. Cannot generate embeddings.")
    exit()


# -------------------------
# TEXT CHUNKING
# -------------------------

chunks = chunk_text(all_text)

print("Total chunks:", len(chunks))


# -------------------------
# EMBEDDING MODEL
# -------------------------

print("Loading embedding model...")

model = SentenceTransformer("BAAI/bge-small-en-v1.5")

embeddings = model.encode(chunks)

print("Embeddings generated.")


# -------------------------
# BUILD FAISS INDEX
# -------------------------

dimension = len(embeddings[0])

index = faiss.IndexFlatL2(dimension)

index.add(np.array(embeddings))

print("Embeddings stored in FAISS index.")


# -------------------------
# QUESTION SEARCH LOOP
# -------------------------

while True:

    query = input("\nAsk a question about the document (type 'exit' to quit): ")

    if query.lower() == "exit":
        break

    query_vector = model.encode([query])

    D, I = index.search(query_vector, k=3)

    print("\nMost relevant sections:\n")

    for idx in I[0]:

        print(chunks[idx])
        print("\n----------------------")